#!/usr/bin/env python3
"""
icmg ingest sidecar — Phase 38.

Protocol (line-delimited JSON over stdin/stdout — same pattern as embedder):
  -> on startup: {"op":"ready","types":["pdf","image","office"]}
  -> requests:
       {"op":"pdf",    "id":N, "path":"..."}    -> {"id":N, "text":"...", "pages":N}
       {"op":"image",  "id":N, "path":"..."}    -> {"id":N, "text":"...", "lang":"eng"}
       {"op":"office", "id":N, "path":"..."}    -> {"id":N, "text":"...", "kind":"docx|xlsx"}
       {"op":"shutdown"}

Graceful: optional deps (pdfplumber, pytesseract, python-docx, openpyxl).
Missing dep -> error per request.
"""
import io
import os
import sys
import json

sys.stdin  = io.TextIOWrapper(sys.stdin.buffer,  encoding="utf-8", errors="replace", newline="\n")
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace", newline="\n", line_buffering=True)
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace", newline="\n")

# Content types this sidecar can extract (advertised in the ready op).
READY_TYPES = ["pdf", "image", "office", "media"]

def emit(obj):
    sys.stdout.write(json.dumps(obj, ensure_ascii=False) + "\n")
    sys.stdout.flush()

def extract_pdf(path):
    try:
        import pdfplumber
    except ImportError:
        return {"error": "pdfplumber not installed (pip install pdfplumber)"}
    try:
        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text_parts.append(page.extract_text() or "")
            n_pages = len(pdf.pages)
        return {"text": "\n\n".join(text_parts), "pages": n_pages}
    except Exception as e:
        return {"error": f"pdf read failed: {e}"}

def extract_image(path):
    try:
        from PIL import Image
        import pytesseract
    except ImportError:
        return {"error": "pytesseract / Pillow not installed (pip install pytesseract Pillow + tesseract binary)"}
    try:
        img = Image.open(path)
        text = pytesseract.image_to_string(img)
        return {"text": text, "lang": "eng"}
    except Exception as e:
        return {"error": f"ocr failed: {e}"}

def extract_office(path):
    """Extract text from an Office document (.docx / .xlsx).

    Pure + graceful: never raises. Returns {"text":..., "kind":...} on success
    or {"error":...} for a missing file, unsupported extension, or absent dep.
    """
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        try:
            import docx
        except ImportError:
            return {"error": "python-docx not installed (pip install python-docx)"}
        try:
            d = docx.Document(path)
            parts = [p.text for p in d.paragraphs if p.text]
            # tables too (rows of tab-joined cells)
            for tbl in d.tables:
                for row in tbl.rows:
                    cells = [c.text for c in row.cells if c.text]
                    if cells:
                        parts.append("\t".join(cells))
            return {"text": "\n".join(parts), "kind": "docx"}
        except Exception as e:
            return {"error": f"docx read failed: {e}"}
    elif ext == ".xlsx":
        try:
            import openpyxl
        except ImportError:
            return {"error": "openpyxl not installed (pip install openpyxl)"}
        try:
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            parts = []
            for ws in wb.worksheets:
                parts.append(f"# {ws.title}")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        parts.append("\t".join(cells))
            wb.close()
            return {"text": "\n".join(parts), "kind": "xlsx"}
        except Exception as e:
            return {"error": f"xlsx read failed: {e}"}
    else:
        return {"error": f"unsupported office extension: {ext or '(none)'} (want .docx/.xlsx)"}

MEDIA_EXTS = {".mp4", ".mov", ".mkv", ".webm", ".avi", ".m4v",
              ".mp3", ".wav", ".m4a", ".flac", ".ogg", ".aac"}

def extract_media(path):
    """Transcribe a video/audio file to text via faster-whisper (G2).

    Pure + graceful: never raises. Returns {"text":..., "kind":"media",
    "lang":...} on success or {"error":...} for a missing file, unsupported
    extension, or absent dependency. faster-whisper bundles its own audio
    decoding (PyAV), so ffmpeg is not required for common containers.
    """
    ext = os.path.splitext(path)[1].lower()
    if ext not in MEDIA_EXTS:
        return {"error": f"unsupported media extension: {ext or '(none)'} "
                         f"(want one of {sorted(MEDIA_EXTS)})"}
    if not os.path.exists(path):
        return {"error": f"file not found: {path}"}
    try:
        from faster_whisper import WhisperModel
    except ImportError:
        return {"error": "faster-whisper not installed "
                         "(pip install faster-whisper)"}
    try:
        # 'base' is a good size/speed tradeoff; CPU int8 keeps it light. The
        # model is cached by faster-whisper after first download.
        model_size = os.environ.get("ICMG_WHISPER_MODEL", "base")
        device = os.environ.get("ICMG_WHISPER_DEVICE", "cpu")
        compute = os.environ.get("ICMG_WHISPER_COMPUTE", "int8")
        model = WhisperModel(model_size, device=device, compute_type=compute)
        segments, info = model.transcribe(path)
        text = " ".join(seg.text.strip() for seg in segments).strip()
        return {"text": text, "kind": "media",
                "lang": getattr(info, "language", "") or ""}
    except Exception as e:
        return {"error": f"media transcription failed: {e}"}

def main():
    emit({"op": "ready", "types": READY_TYPES})
    for line in sys.stdin:
        line = line.strip()
        if not line:
            continue
        try:
            req = json.loads(line)
        except Exception as e:
            emit({"id": 0, "error": f"bad json: {e}"})
            continue
        op = req.get("op", "")
        if op == "shutdown":
            return 0
        rid = req.get("id", 0)
        path = req.get("path", "")
        if op == "pdf":
            r = extract_pdf(path)
        elif op == "image":
            r = extract_image(path)
        elif op == "office":
            r = extract_office(path)
        elif op == "media":
            r = extract_media(path)
        else:
            r = {"error": f"unknown op: {op}"}
        r["id"] = rid
        emit(r)
    return 0

if __name__ == "__main__":
    sys.exit(main())
