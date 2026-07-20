#!/usr/bin/env python3
"""
Standalone tests for the icmg ingest sidecar (multimodal/icmg_ingest.py).

Run:  python multimodal/test_ingest.py
Exit 0 = all pass. No pytest dependency (mirrors the C++ TEST() harness style:
tiny, self-contained, deterministic). Office tests skip gracefully if
python-docx / openpyxl are absent, so CI without those deps stays green.
"""
import os
import sys
import tempfile
import importlib.util

HERE = os.path.dirname(os.path.abspath(__file__))

def load_sidecar():
    spec = importlib.util.spec_from_file_location(
        "icmg_ingest", os.path.join(HERE, "icmg_ingest.py"))
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod

_pass = 0
_fail = 0
def check(name, cond):
    global _pass, _fail
    if cond:
        _pass += 1
        print(f"  [PASS] {name}")
    else:
        _fail += 1
        print(f"  [FAIL] {name}")

def skip(name, why):
    print(f"  [SKIP] {name} ({why})")

def main():
    ing = load_sidecar()

    # --- extract_office must exist ---
    check("sidecar exposes extract_office", hasattr(ing, "extract_office"))
    if not hasattr(ing, "extract_office"):
        print(f"\n{_pass} passed, {_fail} failed")
        return 1 if _fail else 0

    # --- ready op advertises office ---
    # (main() emits ready; we can't easily capture it here without a subprocess,
    #  so just assert the advertised types list is defined and includes office.)
    check("READY_TYPES includes 'office'",
          "office" in getattr(ing, "READY_TYPES", []))

    # --- unknown extension -> graceful error, never throws ---
    r = ing.extract_office("/nonexistent/file.docx")
    check("missing docx returns error dict (no throw)",
          isinstance(r, dict) and "error" in r)

    r = ing.extract_office("/some/file.txt")
    check("unsupported extension returns error dict",
          isinstance(r, dict) and "error" in r)

    # --- real .docx round-trip ---
    try:
        import docx
        p = os.path.join(tempfile.mkdtemp(), "t.docx")
        d = docx.Document()
        d.add_paragraph("Hello from docx")
        d.add_paragraph("Second line here")
        d.save(p)
        r = ing.extract_office(p)
        check("docx: no error", "error" not in r)
        check("docx: text contains paragraph 1", "Hello from docx" in r.get("text", ""))
        check("docx: text contains paragraph 2", "Second line here" in r.get("text", ""))
    except ImportError:
        skip("docx round-trip", "python-docx not installed")

    # --- real .xlsx round-trip ---
    try:
        import openpyxl
        p = os.path.join(tempfile.mkdtemp(), "t.xlsx")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws["A1"] = "Name"
        ws["B1"] = "Score"
        ws["A2"] = "widget"
        ws["B2"] = 42
        wb.save(p)
        r = ing.extract_office(p)
        check("xlsx: no error", "error" not in r)
        txt = r.get("text", "")
        check("xlsx: contains header cell", "Name" in txt and "Score" in txt)
        check("xlsx: contains data cell", "widget" in txt and "42" in txt)
    except ImportError:
        skip("xlsx round-trip", "openpyxl not installed")

    # --- G2: video/audio transcription (faster-whisper) ---
    check("sidecar exposes extract_media", hasattr(ing, "extract_media"))
    check("READY_TYPES includes 'media'",
          "media" in getattr(ing, "READY_TYPES", []))
    if hasattr(ing, "extract_media"):
        # missing file -> graceful error dict, never throws
        r = ing.extract_media("/nonexistent/clip.mp4")
        check("missing media returns error dict (no throw)",
              isinstance(r, dict) and "error" in r)
        # unsupported extension -> graceful error
        r = ing.extract_media("/some/file.txt")
        check("unsupported media extension returns error dict",
              isinstance(r, dict) and "error" in r)
        # dep-absent path: if faster-whisper isn't installed, a real .mp4 name
        # must still return an error dict (not raise). We don't ship a media
        # fixture, so just assert the dep-guard shape on a fake path.
        try:
            import faster_whisper  # noqa: F401
            skip("media transcription", "faster-whisper present; skipping heavy run")
        except ImportError:
            r = ing.extract_media("/tmp/whatever.mp4")
            check("media without faster-whisper -> error dict",
                  isinstance(r, dict) and "error" in r)

    print(f"\n{_pass} passed, {_fail} failed")
    return 1 if _fail else 0

if __name__ == "__main__":
    sys.exit(main())
